from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUT = ROOT / "RawDrive_User_Manual_Luxury.docx"

NAVY = RGBColor(11, 16, 36)
INK = RGBColor(22, 29, 50)
MUTED = RGBColor(92, 102, 125)
CYAN = RGBColor(47, 196, 239)
GOLD = RGBColor(197, 163, 89)
WHITE = RGBColor(255, 255, 255)
PALE = RGBColor(233, 238, 250)

SCREENSHOTS = [
    ("dashboard.png", "Dashboard", "Your daily command center for galleries, storage, quick actions, and recent work."),
    ("galleries.png", "Galleries", "Create, publish, share, and manage all client galleries from one screen."),
    ("gallery-workspace.png", "Gallery workspace", "Upload photos, create sub-galleries, sort assets, and open delivery tools."),
    ("cover-design.png", "Cover and Design", "Set the public gallery look, cover images, grid layout, and folder presentation."),
    ("settings-share.png", "Sharing and settings", "Control access windows, shared accounts, storage billing, and public availability."),
    ("faceid-search.png", "FaceID search", "Sync faces, review detected people, and let clients find themselves by selfie."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{side}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def paragraph(doc: Document, text: str = "", *, style: str | None = None, align=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.16
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=11.5, color=INK)
    return p


def add_label(doc: Document, text: str, after=4):
    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.LEFT, after=after)
    run = p.add_run(text.upper())
    set_run_font(run, size=9.5, color=GOLD, bold=True)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    style_name = f"Heading {level}"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    return p.add_run(text)


def add_image(doc: Document, image_name: str, width_in: float = 6.25) -> None:
    image_path = ASSETS / image_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    p.paragraph_format.space_after = Pt(4)


def add_caption(doc: Document, caption: str) -> None:
    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10)
    run = p.add_run(caption)
    set_run_font(run, size=9.5, color=MUTED, bold=False)


def add_callout(doc: Document, title: str, body: str, fill: str = "EDF6FF") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    title_run = p.add_run(title)
    set_run_font(title_run, size=11.5, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    body_run = p2.add_run(body)
    set_run_font(body_run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_step_table(doc: Document) -> None:
    rows = [
        ("1", "Open dashboard", "Check storage, recent galleries, and quick actions."),
        ("2", "Create gallery", "Choose Delivery for client photos or Proofing for selections."),
        ("3", "Upload photos", "Select files or folders, then keep the upload dashboard open until completion."),
        ("4", "Design public view", "Set cover, mobile cover, grid columns, and folder layout."),
        ("5", "Share with clients", "Publish, copy the gallery link, and set an access window if needed."),
        ("6", "Use FaceID", "Sync gallery photos, review detected people, and let clients search by selfie."),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    set_table_width(table)
    remove_table_borders(table)
    for row, (num, title, detail) in zip(table.rows, rows):
        for cell in row.cells:
            set_cell_shading(cell, "F6F8FC")
            set_cell_margins(cell, top=120, bottom=120, start=120, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].width = Inches(0.45)
        row.cells[1].width = Inches(1.7)
        row.cells[2].width = Inches(4.35)
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(num)
        set_run_font(r0, size=11, color=CYAN, bold=True)
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(title)
        set_run_font(r1, size=10.5, color=NAVY, bold=True)
        p2 = row.cells[2].paragraphs[0]
        r2 = p2.add_run(detail)
        set_run_font(r2, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.36)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        run = p.add_run(f" {item}")
        set_run_font(run, size=10.8, color=INK)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11.5)
    normal.font.color.rgb = INK

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.bold = True
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 1"].font.color.rgb = NAVY
    styles["Heading 2"].font.size = Pt(13.5)
    styles["Heading 2"].font.color.rgb = NAVY
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = MUTED

    if "Luxury Caption" not in styles:
        cap = styles.add_style("Luxury Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Calibri"
        cap.font.size = Pt(9.5)
        cap.font.color.rgb = MUTED

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("RawDrive User Manual | Studio workflow guide")
        set_run_font(footer_run, size=8.5, color=MUTED)


def add_cover(doc: Document) -> None:
    add_label(doc, "RawDrive Studio Guide", after=6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    title = p.add_run("RawDrive User Manual")
    set_run_font(title, size=30, color=NAVY, bold=True)

    p = paragraph(doc, "A simple, visual guide for dashboard, galleries, uploads, sharing, design, and FaceID.", after=14)
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = MUTED

    add_callout(
        doc,
        "Best for studio operators",
        "Use this guide when training team members, setting up a new gallery, or checking the right workflow before client delivery.",
        fill="F9F4E8",
    )
    add_image(doc, "dashboard.png", width_in=6.25)
    add_caption(doc, "Dashboard view used in this manual.")
    paragraph(doc, "Prepared for RawDrive studios | June 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    doc.add_page_break()


def add_section(
    doc: Document,
    label: str,
    title: str,
    image: str | None,
    caption: str | None,
    bullets: list[str],
    callout: tuple[str, str] | None = None,
) -> None:
    add_label(doc, label)
    add_heading(doc, title, 1)
    if image:
        add_image(doc, image)
        if caption:
            add_caption(doc, caption)
    add_bullets(doc, bullets)
    if callout:
        add_callout(doc, callout[0], callout[1])


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_label(doc, "Quick start")
    add_heading(doc, "The complete gallery flow", 1)
    paragraph(
        doc,
        "RawDrive is organized around a simple studio flow. Start on the dashboard, create a gallery, upload photos, style the public view, then share the finished link with clients.",
        after=8,
    )
    add_step_table(doc)
    add_callout(
        doc,
        "Operator rule",
        "Keep uploads running until all files are complete. You can work in other RawDrive screens, but keep the browser session open while large folders finish.",
        fill="EAF8FF",
    )
    doc.add_page_break()

    add_section(
        doc,
        "Dashboard",
        "Read the studio at a glance",
        "dashboard.png",
        "The dashboard shows storage, galleries, quick actions, and recent work.",
        [
            "Use Quick actions for Create gallery, Add client, Send invoice, and New booking.",
            "Watch storage usage before large uploads so client delivery is not blocked.",
            "Open recent galleries directly from the dashboard card.",
        ],
    )

    add_section(
        doc,
        "Galleries",
        "Create and manage client galleries",
        "galleries.png",
        "The Galleries screen is the central list of published and draft galleries.",
        [
            "Click New Gallery to start a client delivery or proofing gallery.",
            "Use the publish toggle only when the gallery is ready for clients.",
            "Use the card menu for actions such as share, cover, phone cover, and delete.",
        ],
    )
    doc.add_page_break()

    add_section(
        doc,
        "Gallery workspace",
        "Upload, organize, and prepare delivery",
        "gallery-workspace.png",
        "The gallery workspace contains upload tools, sub-galleries, and workflow links.",
        [
            "Use Upload Photos for selected files and Upload Folder for complete camera folders.",
            "Create sub-galleries for highlights, rituals, family sets, or client-only collections.",
            "Keep manual photo order when you rearrange images by mouse.",
            "Open workflow links on the right for cover design, FaceID, settings, and delivery.",
        ],
        ("Upload status", "When uploads are active, use the persistent upload panel to watch progress, cancel active uploads, or retry failed items."),
    )

    add_section(
        doc,
        "Cover and design",
        "Style the public gallery",
        "cover-design.png",
        "Cover and Design controls gallery presentation for desktop and mobile.",
        [
            "Set the root gallery desktop and phone cover images from the root folder.",
            "Use folder-specific grid settings when a sub-gallery is selected.",
            "Preview grid columns before saving so the public gallery feels balanced.",
            "Keep the cover safe zone readable on mobile by centering faces and titles.",
        ],
    )
    doc.add_page_break()

    add_section(
        doc,
        "Sharing",
        "Publish, protect, and share access",
        "settings-share.png",
        "Settings lets you publish the gallery, set access windows, and share with another RawDrive account.",
        [
            "Use No expiry only when clients should keep access indefinitely.",
            "Choose 30, 60, 90 days, or a custom date for controlled delivery windows.",
            "Shared accounts can be given access while storage billing stays with the selected workspace.",
            "Ask the recipient to create a RawDrive account before sharing with their email.",
        ],
    )
    doc.add_page_break()

    add_section(
        doc,
        "FaceID",
        "Help clients find themselves",
        "faceid-search.png",
        "FaceID syncs gallery photos and lets clients search by captured selfie.",
        [
            "Click Sync now after upload to index gallery faces.",
            "Review detected people and rename clear identities when useful.",
            "Clients can use Capture and search from the public gallery to find their own photos.",
            "If new photos are uploaded later, run sync again so FaceID includes them.",
        ],
    )
    doc.add_page_break()

    add_label(doc, "Daily checklist")
    add_heading(doc, "Before sharing a finished gallery", 1)
    add_bullets(
        doc,
        [
            "All uploads completed and failed items retried or dismissed.",
            "Gallery photos appear in the expected order.",
            "Cover image and mobile cover are centered and readable.",
            "Grid settings are saved for the selected folder.",
            "Gallery is published only after final review.",
            "Share link opens correctly in client view.",
            "FaceID sync completed for all required folders.",
        ]
    )
    add_callout(
        doc,
        "Support note",
        "If a JPEG upload warns about camera-authored trailer data, the file is still a normal camera JPEG. Retry after the latest upload fix is deployed.",
        fill="FFF7E6",
    )

    doc.core_properties.title = "RawDrive User Manual"
    doc.core_properties.subject = "Simple dashboard and gallery workflow guide"
    doc.core_properties.author = "RawDrive"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
